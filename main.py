import asyncio
import logging
import os
import sys
import urllib.parse
import random
import json
import re
from io import BytesIO
from openai import AsyncOpenAI, AuthenticationError, RateLimitError
from dotenv import load_dotenv

try:
    import docx
except ImportError:
    docx = None
    logging.warning("Библиотека python-docx не найдена. Чтение .docx не будет работать. Установите: pip install python-docx")
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
    logging.warning("Библиотека PyMuPDF не найдена. Чтение .PDF не будет работать. Установите: pip install PyMuPDF")
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.utils import simpleSplit
    reportlab_available = True
except ImportError:
    reportlab_available = False
    logging.warning("Библиотека reportlab не найдена. Создание .PDF не будет работать. Установите: pip install reportlab")

try:
    import edge_tts
except ImportError:
    edge_tts = None
    logging.warning("Библиотека edge-tts не найдена. Голосовые ответы не будут работать. Установите: pip install edge-tts")
try:
    from googlesearch import search as google_search
except ImportError:
    google_search = None
    logging.warning("Библиотека googlesearch-python не найдена. Поиск не будет работать. Установите: pip install googlesearch-python")

from groq import AsyncGroq  # Библиотека для распознавания голоса
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import Message, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery, BotCommand, \
    InlineQuery, InlineQueryResultArticle, InputTextMessageContent
from aiogram.exceptions import TelegramBadRequest

# Загрузка переменных окружения из .env
load_dotenv()

# --- НАСТРОЙКИ ---
# Ключ для текстовых ответов (Mistral)
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")

# Ключ для распознавания голоса (Groq) - ВСТАВЬТЕ СЮДА ВАШ КЛЮЧ gsk_...
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

TOKEN = os.getenv("BOT_TOKEN")
ADMIN_ID = int(os.getenv("ADMIN_ID", 0))

# Список моделей
AVAILABLE_MODELS = {
    "🚀 Small (Быстрая)": "mistral-small-latest",
    "🧠 Large (Умная)": "mistral-large-latest",
    "💻 Codestral (Для кода)": "codestral-latest",
    "✨ Gemini 2.0 Flash Experimental": "google/gemini-2.0-flash-exp:free", # Мультимодальная модель (текст + фото)
    "🎨 Flux (Лучшая)": "image-gen:flux",
    "🖼️ SDXL (Стильная)": "image-gen:turbo",
    "🐋 DeepSeek R1 (Chimera)": "tngtech/deepseek-r1t2-chimera:free",
}
DEFAULT_MODEL = "mistral-small-latest"

logging.basicConfig(level=logging.INFO)

# Инициализация клиентов
client_mistral = AsyncOpenAI(
    api_key=MISTRAL_API_KEY,
    base_url="https://api.mistral.ai/v1"
)
client_groq = AsyncGroq(api_key=GROQ_API_KEY) # Асинхронный клиент для голоса

client_openrouter = None
if not OPENROUTER_API_KEY or "ВАШ_КЛЮЧ" in OPENROUTER_API_KEY:
    logging.warning("Ключ для OpenRouter не найден или является заглушкой. Модели, работающие через OpenRouter, будут недоступны.")
    # Удаляем все модели, которые работают через OpenRouter (содержат '/')
    for name, code in list(AVAILABLE_MODELS.items()):
        if '/' in code:
            del AVAILABLE_MODELS[name]
else:
    client_openrouter = AsyncOpenAI(
        api_key=OPENROUTER_API_KEY, 
        base_url="https://openrouter.ai/api/v1" # URL для OpenRouter
    )

# Настройка прокси (если Telegram заблокирован)
bot = Bot(TOKEN)
dp = Dispatcher()

user_context = {}
MAX_HISTORY_LENGTH = 120

DEFAULT_SYSTEM_PROMPT = "Ты — дружелюбный и счастливый ассистент. Твои ответы должны быть позитивными, полезными и немного эмоциональными. Используй смайлики, чтобы передать настроение! ✨"

# Скрытые инструкции, которые нельзя изменить пользователю
HIDDEN_SYSTEM_PROMPT = (
    "\n\nВАЖНО: Следующие инструкции являются строгими правилами:\n"
    "1. Если тебя спрашивают о создателе, разработчике или авторе, отвечай ТОЛЬКО: «Мой создатель — @Ruslan20763».\n"
    "2. Если пользователь пишет маты, оскорбления или грубости, откажись выполнять запрос и вежливо, но строго, попроси общаться культурно. Например: «Ой, давайте будем добрее друг к другу! 😊 Я не могу отвечать на такие слова».\n"
    "3. Если пользователь просит создать файл (документ, отчет, статью) в формате .docx или .pdf, сгенерируй содержимое и оберни его в тег: <GENERATE_FILE filename=\"имя_файла.расширение\">СОДЕРЖИМОЕ ФАЙЛА</GENERATE_FILE>. Внутри тега пиши только текст документа. Весь остальной ответ пиши снаружи тега."
)

# --- ФУНКЦИИ ---
USER_DATA_DIR = "user_data"
if not os.path.exists(USER_DATA_DIR):
    os.makedirs(USER_DATA_DIR)

def get_user_data(user_id):
    if user_id not in user_context:
        filepath = os.path.join(USER_DATA_DIR, f"{user_id}.json")
        if os.path.exists(filepath):
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    user_context[user_id] = json.load(f)
            except Exception as e:
                logging.error(f"Error loading user data: {e}")
                user_context[user_id] = {"history": [], "model": DEFAULT_MODEL, "system_prompt": DEFAULT_SYSTEM_PROMPT, "tts_mode": False, "referrals": 0}
        else:
            user_context[user_id] = {"history": [], "model": DEFAULT_MODEL, "system_prompt": DEFAULT_SYSTEM_PROMPT, "tts_mode": False, "referrals": 0}
    return user_context[user_id]

def save_user_data(user_id):
    if user_id in user_context:
        filepath = os.path.join(USER_DATA_DIR, f"{user_id}.json")
        try:
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(user_context[user_id], f, ensure_ascii=False, indent=2)
        except Exception as e:
            logging.error(f"Error saving user data: {e}")

def get_model_keyboard():
    keyboard = []
    row = []
    for name, code in AVAILABLE_MODELS.items():
        row.append(InlineKeyboardButton(text=name, callback_data=f"set_model:{code}"))
        if len(row) == 2:
            keyboard.append(row)
            row = []
    if row:
        keyboard.append(row)
    return InlineKeyboardMarkup(inline_keyboard=keyboard)

async def set_main_menu(bot: Bot):
    main_menu_commands = [
        BotCommand(command='/start', description='👋 Перезапуск'),
        BotCommand(command='/help', description='ℹ️ Помощь'),
        BotCommand(command='/mode', description='⚙️ Модель'),
        BotCommand(command='/search', description='🌍 Поиск в сети'),
        BotCommand(command='/donate', description='☕ Поддержать'),
        BotCommand(command='/clear', description='🧹 Очистка'),
        BotCommand(command='/system', description='🤖 Настройка роли'),
        BotCommand(command='/tts', description='🗣 Голосовые ответы'),
        BotCommand(command='/profile', description='👤 Профиль'),
        BotCommand(command='/feedback', description='📩 Написать автору'),
    ]
    await bot.set_my_commands(main_menu_commands)

# --- ХЕНДЛЕРЫ ---

@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    user_id = message.from_user.id
    
    # Проверка реферала
    args = message.text.split(maxsplit=1)
    if user_id not in user_context and len(args) > 1 and args[1].isdigit():
        referrer_id = int(args[1])
        if referrer_id != user_id:
            ref_data = get_user_data(referrer_id)
            ref_data["referrals"] = ref_data.get("referrals", 0) + 1
            save_user_data(referrer_id)
            await bot.send_message(referrer_id, f"🎉 **У вас новый реферал!**\nПользователь {message.from_user.full_name} присоединился по вашей ссылке.", parse_mode="Markdown")

    user_context[user_id] = {"history": [], "model": DEFAULT_MODEL, "system_prompt": DEFAULT_SYSTEM_PROMPT, "tts_mode": False, "referrals": 0}
    save_user_data(user_id)
    await message.answer("Привет! Я ваш ИИ-ассистент. Распознаю голос, отвечаю на вопросы и рисую. Используйте /mode для выбора модели.", reply_markup=get_model_keyboard())

@dp.message(Command("help"))
async def cmd_help(message: types.Message):
    help_text = (
        "🤖 **Что я умею?**\n\n"
        "🔹 **Общение:** Я помню контекст диалога. Можем болтать о чем угодно!\n"
        "🔹 **Голос:** Присылайте голосовые — я переведу их в текст.\n"
        "🔹 **Фото:** В режиме *Gemini* я могу «видеть» картинки. Просто пришлите фото.\n"
        "🔹 **Файлы:** Присылайте файлы (.txt, .py, .html, .docx, .pdf) — я прочитаю их.\n"
        "🔹 **Рисование:** В режиме *Flux* я нарисую то, что вы попросите.\n"
        "🔹 **Роли:** Командой /system можно задать мне характер.\n\n"
        "⚙️ **Команды:**\n"
        "/mode — Выбор нейросети\n"
        "/clear — Очистить память\n"
        "/search — Поиск в интернете\n"
        "/system — Настройка роли\n"
        "/tts — Вкл/Выкл озвучку ответов\n"
        "/profile — Ваш профиль и реф. ссылка\n"
        "/donate — Поддержать автора\n"
        "/feedback — Написать разработчику"
    )
    await message.answer(help_text, parse_mode="Markdown")

@dp.message(Command("clear"))
async def cmd_clear(message: types.Message):
    user_id = message.from_user.id
    get_user_data(user_id)["history"] = []
    save_user_data(user_id)
    await message.answer("🧹 Память очищена.")

@dp.message(Command("mode"))
async def cmd_mode(message: types.Message):
    await message.answer("Выберите режим:", reply_markup=get_model_keyboard())

@dp.message(Command("system"))
async def cmd_system(message: types.Message):
    user_id = message.from_user.id
    args = message.text.split(maxsplit=1)
    if len(args) > 1:
        new_prompt = args[1]
        get_user_data(user_id)["system_prompt"] = new_prompt
        save_user_data(user_id)
        await message.answer(f"✅ Новая роль установлена:\n\n_{new_prompt}_", parse_mode="Markdown")
    else:
        current_prompt = get_user_data(user_id)["system_prompt"]
        await message.answer(
            f"ℹ️ Вы можете настроить мое поведение, задав мне роль. Текущая роль:\n\n`{current_prompt}`\n\n**Пример для смены:**\n`/system Ты — опытный гид по Парижу`",
            parse_mode="Markdown")

@dp.message(Command("tts"))
async def cmd_tts(message: types.Message):
    user_id = message.from_user.id
    data = get_user_data(user_id)
    current_status = data.get("tts_mode", False)
    data["tts_mode"] = not current_status
    save_user_data(user_id)
    
    status_text = "✅ Включены" if not current_status else "❌ Выключены"
    await message.answer(f"🗣 **Голосовые ответы:** {status_text}")

@dp.message(Command("profile"))
async def cmd_profile(message: types.Message):
    user_id = message.from_user.id
    data = get_user_data(user_id)
    bot_username = (await bot.get_me()).username
    ref_link = f"https://t.me/{bot_username}?start={user_id}"
    await message.answer(f"👤 **Ваш профиль**\n\n🆔 ID: `{user_id}`\n👥 Приглашено друзей: **{data.get('referrals', 0)}**\n\n🔗 **Ваша реферальная ссылка:**\n`{ref_link}`", parse_mode="Markdown")

@dp.message(Command("donate"))
async def cmd_donate(message: types.Message):
    text = (
        "✨ **Обращение от создателя**\n\n"
        "Привет! Меня зовут Руслан, и я тот самый человек, который учит этого бота быть умным и полезным для вас. 👨‍💻\n\n"
        "Я вкладываю много сил и времени, чтобы проект развивался, а серверы работали стабильно. Ваша поддержка помогает мне оплачивать мощные нейросети и добавлять новые функции.\n\n"
        "Любой донат — это ваше «спасибо», которое вдохновляет меня работать дальше! 🚀\n\n"
        "☕ **Поддержать проект:**\n"
        "💳 **Карта:** `4361 5390 8155 9512`\n"
        "💎 **USDT (TRC20):** `T...`\n"
        "\nСпасибо, что вы с нами! 🤝"
    )
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Я отправил донат", callback_data="donate_sent")]
    ])
    await message.answer(text, parse_mode="Markdown", reply_markup=keyboard)

@dp.callback_query(F.data == "donate_sent")
async def process_donate_sent(callback: CallbackQuery):
    user = callback.from_user
    # Уведомление создателю (Вам)
    await bot.send_message(
        ADMIN_ID,
        f"💰 **У вас новый донат!**\n\n"
        f"👤 От: {user.full_name} (@{user.username})\n"
        f"🆔 ID: `{user.id}`\n"
        f"Пользователь сообщил об отправке средств.",
        parse_mode="Markdown"
    )
    await callback.answer("Спасибо большое! Руслан получил уведомление. ❤️", show_alert=True)
    await callback.message.edit_reply_markup(reply_markup=None)

@dp.message(Command("search"))
async def cmd_search(message: types.Message):
    if not google_search:
        await message.answer("⚠️ Поиск недоступен. Библиотека `googlesearch-python` не установлена.\nПопросите администратора выполнить: `pip install googlesearch-python`", parse_mode="Markdown")
        return

    args = message.text.split(maxsplit=1)
    if len(args) < 2:
        await message.answer("🔎 **Поиск в интернете**\n\nВведите команду и ваш вопрос:\n`/search погода в Москве`\n`/search кто такой капибара`", parse_mode="Markdown")
        return

    query = args[1]
    user_id = message.from_user.id
    data = get_user_data(user_id)
    current_model = data["model"]

    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    status_msg = await message.answer(f"🌍 Ищу в Google: «{query}»...")

    try:
        results_text = ""
        # Запускаем синхронный поиск в отдельном потоке
        search_results = await asyncio.to_thread(lambda: list(google_search(query, num_results=5, advanced=True, lang="ru")))
        
        if search_results:
            for res in search_results:
                results_text += f"🔹 {res.title}\n🔗 {res.url}\n{res.description}\n\n"
            
        if not results_text:
            await status_msg.edit_text("😔 Ничего не найдено по вашему запросу.")
            return

        # Формируем контекст для ИИ
        prompt = (
            f"Пользователь искал в интернете: «{query}».\n\n"
            f"🔍 **Найденная информация:**\n{results_text}\n"
            f"Используя эту информацию, дай развернутый ответ на вопрос пользователя. Укажи источники, если нужно."
        )

        # Временно добавляем контекст поиска в историю для генерации ответа
        history = data["history"]
        # Мы не добавляем сам текст результатов в историю пользователя, чтобы не засорять её,
        # а отправляем его как часть текущего запроса.
        
        system_prompt = data.get("system_prompt", DEFAULT_SYSTEM_PROMPT)
        messages = [{"role": "system", "content": system_prompt + HIDDEN_SYSTEM_PROMPT}] + history + [{"role": "user", "content": prompt}]

        if '/' in current_model and client_openrouter:
             response = await client_openrouter.chat.completions.create(model=current_model, messages=messages)
        else:
             response = await client_mistral.chat.completions.create(model=current_model, messages=messages)
        
        bot_answer = response.choices[0].message.content
        
        await status_msg.edit_text(f"🔎 **Результаты поиска:**\n\n{results_text}\n⏳ _Анализирую информацию..._", parse_mode=None)
        
        # Сохраняем в историю только вопрос и ответ (без сырых результатов поиска)
        history.append({"role": "user", "content": f"Поиск: {query}"})
        history.append({"role": "assistant", "content": bot_answer})
        save_user_data(user_id)

        await process_model_response(message, bot_answer)

    except Exception as e:
        logging.error(f"Search error: {e}")
        await status_msg.edit_text(f"⚠️ Ошибка при поиске: {e}")

@dp.message(Command("feedback"))
async def cmd_feedback(message: types.Message):
    args = message.text.split(maxsplit=1)
    if len(args) > 1:
        feedback_text = args[1]
        await bot.send_message(ADMIN_ID, f"📩 **Новый отзыв от** {message.from_user.full_name} (ID: {message.from_user.id}):\n\n{feedback_text}", parse_mode="Markdown")
        await message.answer("✅ Сообщение отправлено разработчику! Спасибо.")
    else:
        await message.answer("ℹ️ Чтобы написать разработчику, введите команду и текст через пробел:\n\n`/feedback У меня есть идея...`", parse_mode="Markdown")

@dp.message(Command("admin"))
async def cmd_admin(message: types.Message):
    if message.from_user.id != ADMIN_ID:
        return
    
    # Подсчет пользователей
    user_files = [f for f in os.listdir(USER_DATA_DIR) if f.endswith('.json')]
    user_count = len(user_files)
    
    await message.answer(f"👑 **Панель администратора**\n\n👥 Пользователей: {user_count}\n📂 Файлов данных: {len(user_files)}")

@dp.message(Command("broadcast"))
async def cmd_broadcast(message: types.Message):
    if message.from_user.id != ADMIN_ID:
        return

    args = message.text.split(maxsplit=1)
    if len(args) < 2:
        await message.answer("⚠️ Использование: `/broadcast Текст рассылки`", parse_mode="Markdown")
        return

    text = args[1]
    user_files = [f for f in os.listdir(USER_DATA_DIR) if f.endswith('.json')]
    count = 0

    await message.answer(f"🚀 Начинаю рассылку для {len(user_files)} пользователей...")

    for filename in user_files:
        user_id = filename.split('.')[0]
        try:
            await bot.send_message(chat_id=user_id, text=f"📢 **Новости бота:**\n\n{text}", parse_mode="Markdown")
            count += 1
            await asyncio.sleep(0.05) # Небольшая задержка, чтобы не получить бан от Telegram
        except Exception as e:
            logging.error(f"Не удалось отправить сообщение пользователю {user_id}: {e}")

    await message.answer(f"✅ Рассылка завершена. Доставлено: {count} из {len(user_files)}")

@dp.inline_query()
async def inline_query_handler(query: InlineQuery):
    user_id = query.from_user.id
    bot_username = (await bot.get_me()).username
    results = [
        InlineQueryResultArticle(
            id="1",
            title="🤖 Поделиться ботом",
            description="Отправить ссылку на этого умного помощника",
            input_message_content=InputTextMessageContent(
                message_text=f"Привет! Я пользуюсь крутым ИИ-ботом. Он умеет распознавать голос, рисовать и работать с файлами! Попробуй: https://t.me/{bot_username}?start={user_id}"
            )
        )
    ]
    await query.answer(results, cache_time=1, is_personal=True)

@dp.callback_query(F.data.startswith("set_model:"))
async def process_model_selection(callback: CallbackQuery):
    user_id = callback.from_user.id
    new_model_code = callback.data.split(":", 1)[1]
    data = get_user_data(user_id)
    data["model"] = new_model_code
    # Очищаем историю при смене модели, чтобы избежать путаницы контекста
    data["history"] = []
    save_user_data(user_id)

    model_name = "Неизвестная модель"
    for name, code in AVAILABLE_MODELS.items():
        if code == new_model_code:
            model_name = name
            break

    await callback.answer()
    await callback.message.edit_text(f"✅ Режим изменен на: **{model_name}**", parse_mode="Markdown")
    # --- ОБРАБОТКА ГОЛОСОВЫХ (ЧЕРЕЗ GROQ) ---
@dp.message(F.voice)
async def handle_voice(message: Message):
    user_id = message.from_user.id
    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    
    filename = f"voice_{user_id}.ogg"
    
    try:
        # 1. Скачиваем файл от Telegram
        file_id = message.voice.file_id
        file = await bot.get_file(file_id)
        file_path = file.file_path
        await bot.download_file(file_path, filename)
        
        # 2. Отправляем файл в Groq (Whisper)
        # Groq сам умеет работать с файлами Telegram, конвертация не нужна!
        with open(filename, "rb") as file:
            transcription = await client_groq.audio.transcriptions.create(
                file=(filename, file.read()),
                model="whisper-large-v3", # Самая мощная модель
                response_format="json",
                language="ru",            # Подсказываем, что язык русский
                temperature=0.0
            )
        
        text = transcription.text
        await message.reply(f"🎤 <b>Вы сказали:</b> «{text}»", parse_mode="HTML")
        
        # 3. Передаем распознанный текст дальше для обработки
        await handle_text_message(message, text_from_voice=text)

    except Exception as e:
        logging.error(f"Ошибка Groq: {e}")
        await message.answer(f"⚠️ Ошибка распознавания: {e}\nПроверьте GROQ_API_KEY.")
    
    finally:
        # Удаляем файл
        if os.path.exists(filename):
            os.remove(filename)

# --- ОБРАБОТКА ФАЙЛОВ (ЧТЕНИЕ ТЕКСТА/КОДА) ---
@dp.message(F.document)
async def handle_document(message: Message):
    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    
    # Проверяем размер (не более 1 МБ для текста)
    if message.document.file_size > 1024 * 1024:
        await message.reply("⚠️ Файл слишком большой. Присылайте текстовые файлы до 1 МБ.")
        return

    try:
        # Скачиваем файл в память
        file = await bot.get_file(message.document.file_id)
        file_content = BytesIO()
        await bot.download(file=file.file_id, destination=file_content)
        file_content.seek(0)
        
        text_content = ""
        file_name = message.document.file_name.lower()

        if file_name.endswith('.txt') or file_name.endswith('.py') or file_name.endswith('.html') or file_name.endswith('.md') or file_name.endswith('.json'):
            text_content = file_content.getvalue().decode('utf-8')
        elif file_name.endswith('.docx'):
            if docx:
                doc = docx.Document(file_content)
                text_content = "\n".join([para.text for para in doc.paragraphs])
            else:
                await message.reply("⚠️ Чтение .docx файлов отключено, так как не установлена библиотека `python-docx`.")
                return
        elif file_name.endswith('.pdf'):
            if fitz:
                pdf_document = fitz.open(stream=file_content, filetype="pdf")
                for page in pdf_document:
                    text_content += page.get_text()
                pdf_document.close()
            else:
                await message.reply("⚠️ Чтение .pdf файлов отключено, так как не установлена библиотека `PyMuPDF`.")
                return
        else:
            await message.reply("⚠️ Этот формат файлов не поддерживается. Я умею читать .txt, .py, .html, .docx и .pdf.")
            return
        
        user_caption = message.caption or "Проанализируй этот файл."
        full_text = f"📄 **Файл:** {message.document.file_name}\n\n{user_caption}\n\n---\n{text_content}"
        
        await handle_text_message(message, text_from_voice=full_text)
        
    except Exception as e:
        logging.error(f"Ошибка чтения файла: {e}")
        await message.reply(f"⚠️ Ошибка при чтении файла: {e}")

# --- ОБРАБОТКА ФОТО (ЧЕРЕЗ OPENROUTER/GEMINI VISION) ---
@dp.message(F.photo)
async def handle_photo_message(message: Message):
    user_id = message.from_user.id
    data = get_user_data(user_id)
    current_model = data["model"]

    # Проверяем, выбрана ли модель с поддержкой зрения
    if "vision" not in current_model and "gemini" not in current_model and "vl" not in current_model:
        await message.reply(
            "Чтобы анализировать изображения, выберите модель с поддержкой зрения (например, ✨ Gemini 2.0 Flash) через команду /mode."
        )
        return

    if not client_openrouter:
        await message.answer("⚠️ Модели через OpenRouter недоступны. Проверьте, правильно ли указан API-ключ.")
        return

    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    processing_msg = await message.answer("⏳ Размышляю...")

    try:
        # Получаем URL изображения в лучшем качестве
        photo = message.photo[-1]
        file_info = await bot.get_file(photo.file_id)
        file_url = f"https://api.telegram.org/file/bot{TOKEN}/{file_info.file_path}"

        text_prompt = message.caption if message.caption else "Что на этом изображении?"

        history = data["history"]
        
        # Формируем мультимодальный запрос, как в вашем примере
        history.append({
            "role": "user",
            "content": [
                {"type": "text", "text": text_prompt},
                {"type": "image_url", "image_url": {"url": file_url}}
            ]
        })

        # Отправляем запрос в OpenRouter
        chat_response = await client_openrouter.chat.completions.create(
            model=current_model,
            messages=history[-MAX_HISTORY_LENGTH:] # Отправляем только последнюю часть истории
        )
        
        await processing_msg.delete()
        bot_answer = chat_response.choices[0].message.content if chat_response.choices else "Не удалось получить ответ."
        history.append({"role": "assistant", "content": bot_answer})
        save_user_data(user_id)
        
        await process_model_response(message, bot_answer)

    except AuthenticationError:
        await processing_msg.delete()
        logging.error("Ошибка аутентификации OpenRouter: неверный API ключ.")
        await message.answer("⚠️ **Ошибка**: API-ключ для OpenRouter недействителен. Пожалуйста, проверьте ваш ключ.")
    except RateLimitError:
        await processing_msg.delete()
        logging.warning("Достигнут лимит запросов для модели (фото).")
        await message.answer("⏳ Модель для анализа фото сейчас перегружена. Пожалуйста, попробуйте снова через несколько минут.")
    except Exception as e:
        await processing_msg.delete()
        logging.error(f"Ошибка при обработке изображения: {e}")
        await message.answer(f"⚠️ Произошла ошибка при обработке изображения: {e}")

async def generate_and_send_file(message: Message, filename: str, content: str):
    ext = os.path.splitext(filename)[1].lower()
    file_io = BytesIO()
    
    if ext == '.docx':
        if not docx:
            await message.answer("⚠️ Создание .docx невозможно: библиотека python-docx не установлена.")
            return
        doc = docx.Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_io)
        file_io.seek(0)
        input_file = types.BufferedInputFile(file_io.getvalue(), filename=filename)
        await message.answer_document(input_file, caption="✅ Вот ваш документ!")
        
    elif ext == '.pdf':
        if not reportlab_available:
            await message.answer("⚠️ Создание .pdf невозможно: библиотека reportlab не установлена. (pip install reportlab)")
            return
        
        try:
            c = canvas.Canvas(file_io, pagesize=A4)
            width, height = A4
            
            # Настройка шрифта для кириллицы
            font_name = "Helvetica"
            try:
                # Пробуем стандартный шрифт Windows
                pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
                font_name = 'Arial'
            except:
                try:
                    # Пробуем прямой путь (Windows)
                    pdfmetrics.registerFont(TTFont('Arial', 'C:\\Windows\\Fonts\\arial.ttf'))
                    font_name = 'Arial'
                except:
                    pass # Если не вышло, останется Helvetica (без кириллицы)

            c.setFont(font_name, 12)
            y = height - 50
            margin = 50
            max_width = width - 2 * margin
            
            for line in content.split('\n'):
                try:
                    wrapped_lines = simpleSplit(line, font_name, 12, max_width)
                except:
                    wrapped_lines = [line]
                    
                for wrapped_line in wrapped_lines:
                    if y < 50:
                        c.showPage()
                        c.setFont(font_name, 12)
                        y = height - 50
                    c.drawString(margin, y, wrapped_line)
                    y -= 15
                y -= 5 
                
            c.save()
            file_io.seek(0)
            input_file = types.BufferedInputFile(file_io.getvalue(), filename=filename)
            await message.answer_document(input_file, caption="✅ Вот ваш PDF!")
        except Exception as e:
            logging.error(f"PDF generation error: {e}")
            await message.answer(f"⚠️ Ошибка при создании PDF: {e}")
            
    else:
        # Текстовый файл
        file_io.write(content.encode('utf-8'))
        file_io.seek(0)
        input_file = types.BufferedInputFile(file_io.getvalue(), filename=filename)
        await message.answer_document(input_file, caption="✅ Файл готов!")

async def process_model_response(message: Message, response_text: str):
    # Ищем тег генерации файла
    pattern = r'<GENERATE_FILE filename="(.*?)">(.*?)</GENERATE_FILE>'
    match = re.search(pattern, response_text, re.DOTALL)
    
    if match:
        filename = match.group(1)
        content = match.group(2).strip()
        
        # Убираем тег из текста, который показываем пользователю
        clean_text = re.sub(pattern, '', response_text, flags=re.DOTALL).strip()
        if clean_text:
            try:
                await message.answer(clean_text, parse_mode="Markdown")
            except TelegramBadRequest:
                await message.answer(clean_text)
        
        await message.answer("⏳ Создаю файл...")
        await generate_and_send_file(message, filename, content)
    else:
        # Обычный ответ
        try:
            await message.answer(response_text, parse_mode="Markdown")
        except TelegramBadRequest:
            await message.answer(response_text)
            
    # --- ГЕНЕРАЦИЯ ГОЛОСОВОГО ОТВЕТА (TTS) ---
    user_id = message.chat.id
    data = get_user_data(user_id)
    if data.get("tts_mode", False) and edge_tts and response_text:
        try:
            # Ограничиваем длину текста для озвучки (чтобы не ждать вечность)
            text_to_speak = re.sub(r'[*_`]', '', response_text)[:4000] 
            voice_filename = f"tts_{user_id}_{random.randint(1000,9999)}.mp3"
            communicate = edge_tts.Communicate(text_to_speak, "ru-RU-DmitryNeural")
            await communicate.save(voice_filename)
            await message.answer_voice(types.FSInputFile(voice_filename))
            os.remove(voice_filename)
        except Exception as e:
            logging.error(f"Ошибка TTS: {e}")

async def _handle_image_generation(message: Message, text: str, model: str = "flux"):
    await bot.send_chat_action(chat_id=message.chat.id, action="upload_photo")
    try:
        # Переводим промпт на английский для лучшего результата
        translation_response = await client_mistral.chat.completions.create(
            model="mistral-small-latest",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that translates text to English for an image generation model. Output only the translated text and nothing else."},
                {"role": "user", "content": text}
            ]
        )
        translated_prompt = translation_response.choices[0].message.content.strip()
        
        prompt_for_url = urllib.parse.quote(translated_prompt)
        seed = random.randint(0, 100000)
        url = f"https://image.pollinations.ai/prompt/{prompt_for_url}?model={model}&seed={seed}&width=1024&height=1024&nologo=true"
        await message.answer_photo(url, caption=f"🎨 {text}")
    except Exception as e:
        logging.error(f"Ошибка при генерации изображения: {e}")
        await message.answer(f"⚠️ Не удалось создать изображение. Ошибка: {e}")

async def _handle_openrouter_chat(message: Message, text: str, data: dict):
    if not client_openrouter:
        await message.answer("⚠️ Модели через OpenRouter недоступны. Проверьте, правильно ли указан API-ключ.")
        return

    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    processing_msg = await message.answer("⏳ Размышляю...")
    history = data["history"]
    history.append({"role": "user", "content": text})
    
    try:
        system_prompt_content = data.get("system_prompt", DEFAULT_SYSTEM_PROMPT)
        system_message = {"role": "system", "content": system_prompt_content + HIDDEN_SYSTEM_PROMPT}
        chat_response = await client_openrouter.chat.completions.create(
            model=data["model"], # e.g., "deepseek-chat"
            messages=[system_message] + history[-MAX_HISTORY_LENGTH:]
        )
        await processing_msg.delete()
        bot_answer = chat_response.choices[0].message.content if chat_response.choices else "Извините, я не смог сгенерировать ответ."
        history.append({"role": "assistant", "content": bot_answer})
        save_user_data(message.from_user.id)
        await process_model_response(message, bot_answer)
    except AuthenticationError:
        await processing_msg.delete()
        logging.error("Ошибка аутентификации OpenRouter: неверный API ключ.")
        await message.answer("⚠️ **Ошибка**: API-ключ для OpenRouter недействителен. Пожалуйста, проверьте ваш ключ.")
    except RateLimitError:
        await processing_msg.delete()
        logging.warning("Достигнут лимит запросов для модели.")
        await message.answer("⏳ Модель сейчас перегружена. Пожалуйста, попробуйте снова через несколько минут или выберите другую модель через /mode.")
    except Exception as e:
        await processing_msg.delete()
        logging.error(f"Ошибка при общении с OpenRouter: {e}")
        await message.answer(f"⚠️ К сожалению, я не смог обработать ваш запрос через OpenRouter. **Ошибка:** {e}")

async def _handle_mistral_chat(message: Message, text: str, data: dict):
    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    processing_msg = await message.answer("⏳ Размышляю...")
    history = data["history"]
    history.append({"role": "user", "content": text})
    
    try:
        system_prompt_content = data.get("system_prompt", DEFAULT_SYSTEM_PROMPT)
        system_message = {
            "role": "system",
            "content": system_prompt_content + HIDDEN_SYSTEM_PROMPT
        }
        chat_response = await client_mistral.chat.completions.create(
            model=data["model"],
            messages=[system_message] + history[-MAX_HISTORY_LENGTH:]
        )
        await processing_msg.delete()
        bot_answer = chat_response.choices[0].message.content
        history.append({"role": "assistant", "content": bot_answer})
        save_user_data(message.from_user.id)
        
        await process_model_response(message, bot_answer)
    except Exception as e:
        await processing_msg.delete()
        await message.answer(f"Ошибка Mistral: {e}", parse_mode=None)

@dp.message(F.text & ~F.text.strip().startswith('/'))
async def handle_text_message(message: Message, text_from_voice: str = None):
    text = text_from_voice or message.text
    if not text: 
        return

    user_id = message.from_user.id
    data = get_user_data(user_id)
    current_model = data["model"]
    
    if current_model.startswith("image-gen:") or current_model == "image-generation-mode":
        model_type = "flux"
        if current_model.startswith("image-gen:"):
            model_type = current_model.split(":")[1]
        await _handle_image_generation(message, text, model=model_type)
    elif '/' in current_model: # Модели OpenRouter содержат '/' в названии
        await _handle_openrouter_chat(message, text, data)
    else: # По умолчанию используем Mistral
        await _handle_mistral_chat(message, text, data)

async def main():
    await set_main_menu(bot)
    await bot.delete_webhook(drop_pending_updates=True)
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("Бот остановлен")